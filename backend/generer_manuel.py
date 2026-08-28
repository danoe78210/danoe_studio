#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_manuel.py – Construit le manuel d'utilisation Word (avec screenshots).
Les images sont lues dans Captures/ ; un encadré gris est posé si l'une manque.
Dépendance : pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
CAPTURES = os.path.join(BASE, 'Captures')
SORTIE = os.path.join(BASE, 'Manuel_utilisation.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def p(t, gras=False, italique=False):
    r = doc.add_paragraph().add_run(t)
    r.bold, r.italic = gras, italique

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def tableau(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, htext in enumerate(headers):
        t.rows[0].cells[i].text = htext
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

def capture(nom, legende):
    chemin = os.path.join(CAPTURES, nom)
    if os.path.isfile(chemin):
        doc.add_picture(chemin, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(legende)
        r.italic = True
        r.font.size = Pt(9)
    else:
        r = doc.add_paragraph().add_run(
            f'[ Screenshot à insérer : {legende} — déposer {nom} dans Captures/ ]')
        r.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Page de titre ──
t = doc.add_heading('Manuel d’utilisation — Danoë Studio', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Générateur universel de roman · KDP — écrit pour les débutants', italique=True)
p('Interface v1.20 · Moteur v2.7', italique=True)

h1('Glossaire')
tableau(['Mot', 'Signification'], [
    ['KDP', 'Service d’auto-publication d’Amazon.'],
    ['.docx', 'Fichier Word, format source du livre.'],
    ['.md', 'Fichier texte simple des chapitres.'],
    ['Terminal', 'Fenêtre noire où l’on tape les commandes.'],
    ['Onglet Excel', 'Feuillet en bas du fichier Excel.'],
])

# ── 1. Fichiers ──
h1('1. Liste des fichiers et descriptions')
tableau(['Fichier / dossier', 'À quoi ça sert ?', 'Y toucher ?'], [
    ['interface_livre.py', 'Fenêtre du studio : boutons, console, statistiques.', 'Non'],
    ['generer_roman.py', 'Moteur qui fabrique le livre.', 'Non'],
    ['IA_Roman.py', 'Résumés et traduction IA.', 'Non'],
    ['coherence_roman.py', 'Contrôle de continuité.', 'Non'],
    ['Configuration_roman.xlsx', 'Tableau de bord du roman.', 'OUI'],
    ['Chapitres/', 'Textes des chapitres (.md).', 'OUI'],
    ['Images/', 'Illustrations.', 'OUI'],
    ['Traductions/', 'Fichiers traduits en anglais.', 'Non'],
    ['_cache_HD/', 'Cache des images 300 dpi.', 'Non'],
    ['*_KDP.docx / *_KDP.pdf', 'Livre final Word / PDF.', 'OUI'],
])

# ── 2. Installation ──
h1('2. Installation et utilisation')
h2('2.1 Installer les outils nécessaires')
p('Étape 1 — Python : téléchargez-le sur python.org, cochez « Add Python to PATH », puis « Install Now ».')
p('Étape 2 — Bibliothèques, dans un terminal :')
p('pip install customtkinter python-docx openpyxl pywin32 Pillow', gras=True)
bullets([
    'customtkinter : dessine la fenêtre.',
    'python-docx : écrit le Word.',
    'openpyxl : lit/écrit l’Excel.',
    'pywin32 : pilote Word (pages, TDM, PDF).',
    'Pillow : images noir & blanc 300 dpi.',
])
p('Étape 3 — Word (indispensable) et Excel. Étape 4 — IA facultative (Ollama / OpenAI, onglet IA).')

h2('2.2 Ouvrir un terminal et lancer les commandes')
bullets([
    'Méthode A : dans l’Explorateur, barre d’adresse → tapez cmd → Entrée.',
    'Méthode B : Maj + clic droit dans le dossier → « Ouvrir la fenêtre PowerShell ici ».',
    'Méthode C : Démarrer → cmd → cd "C:\\votre\\dossier".',
])
tableau(['Commande', 'Effet'], [
    ['python interface_livre.py', 'Ouvre la fenêtre du studio (recommandé).'],
    ['python generer_roman.py', 'Génère le livre sans fenêtre.'],
    ['python generer_roman.py --init', 'Crée seulement l’Excel.'],
    ['python generer_roman.py --verif', 'Vérifie les formats.'],
    ['python generer_manuel.py', 'Fabrique ce manuel Word.'],
])
capture('01_interface_principale.png', 'Fig. 1 — La fenêtre au lancement, console « prêt ».')

# ── 3. Fonctionnalités ──
h1('3. Fonctionnalités')
h2('3.1 L’interface principale, bouton par bouton')
capture('02_barre_laterale.png', 'Fig. 2 — La barre latérale et ses quatre sections.')
h3('Production')
bullets([
    '▶ Générer le livre : construit le Word complet ; boutons grisés + bandeau K2000 pendant le travail.',
    '🖨 PDF KDP noir & blanc : convertit le dernier .docx en PDF imprimerie.',
    '📄 Ouvrir le .docx : ouvre le dernier livre dans Word.',
])
h3('Configuration / Intelligence / Traduction')
bullets([
    '📋 Excel Informations · 📚 Excel Chapitres · 🖼 Images · 📁 Chapitres.',
    '🤖 Résumés IA · 🔍 Cohérence IA.',
    'Liste déroulante (dossier ou chapitre) · ↻ rafraîchir · 🌐 Traduction EN.',
])
p('Cartes du haut : Ouvrage, Format, Mots, Pages, Chapitres, Illustrations. '
  'Console colorée (vert = succès, rouge = alerte, doré = étape). '
  'Barre d’état : message + durée. Infobulles au survol.')
capture('03_console_k2000.png', 'Fig. 3 — Console et bandeau K2000 pendant une génération.')
capture('04_cartes_statistiques.png', 'Fig. 4 — Cartes remplies après génération.')

h2('3.2 Toutes les règles de la génération d’un livre')
h3('Sources, liminaires, structure')
bullets([
    'Lit les onglets Informations / Chapitres (ordre haut→bas) / Style, plus Chapitres/ et Images/.',
    'Liminaires fixes : titre, copyright, puis avertissement, dédicace, épigraphe si remplis.',
    'Actes sur page entière (une ligne, non numérotée) ; chapitres sur page impaire ; '
    'sous-chapitres sur nouvelle page ; séparateurs « --- ✦ --- » ; TDM en fin.',
])
h3('Typographie et pagination')
bullets([
    'Corps justifié (Aptos 11 par défaut), titres Cinzel, lettrine +3 pt, 1er paragraphe indenté 0,5 cm.',
    'Gras/italique Markdown conservés.',
    'Départs sur page impaire, page blanche sans numéro si besoin, numéro 1 au premier chapitre, marges miroir.',
])
h3('Marges KDP')
tableau(['Pages', 'Gouttière', 'Extérieure'], [
    ['≤ 150', '0,375 po', '0,250 po'], ['≤ 300', '0,5 po', '0,313 po'],
    ['≤ 500', '0,625 po', '0,375 po'], ['≤ 700', '0,75 po', '0,5 po'],
    ['≤ 828', '0,875 po', '0,625 po'],
])
p('Haut/bas 1,9 cm ; en-têtes/pieds 1 cm. Marges figées avant construction.')
h3('En-têtes, pieds, illustrations, corrections, statistiques')
capture('05_excel_chapitres.png', 'Fig. 5 — L’onglet Chapitres qui pilote l’ordre du livre.')
bullets([
    'En-tête : titre du chapitre en italique souligné ; pas d’en-tête en 1ʳ page de chapitre ; pied numéroté.',
    'Images : ordre Excel, pleine page centrée, N&B 300 dpi, taille adaptée, légende 9 pt ; repère gris si absente.',
    'Corrections : apostrophes ’, pensées en « », espaces doubles, coquilles connues, caractères invisibles.',
    'Statistiques : mots, pages réelles, chapitres, illustrations, lecture à 230 mots/min, détail par chapitre.',
])

h2('3.3 Toutes les règles de la génération du PDF')
capture('06_export_pdf.png', 'Fig. 6 — Le statut après un export PDF réussi.')
bullets([
    'Reprend le dernier *_KDP.docx sur une copie (original préservé).',
    'Texte forcé en noir pur ; images non compressées 300 dpi.',
    'Export Word qualité impression ; propriétés + balises de structure incluses.',
    'PDF non ouvert automatiquement ; contrôle de taille ±2 pt ; message clair si Word/pywin32 absent.',
])

h1('Annexe — Dépannage express')
tableau(['Symptôme', 'Solution'], [
    ['« pip n’est pas reconnu »', 'Réinstaller Python en cochant « Add Python to PATH ».'],
    ['« requiert customtkinter »', 'Lancer la commande pip install de l’étape 2.'],
    ['« Aucun .docx généré »', 'Cliquer d’abord sur ▶ Générer le livre.'],
    ['Ligne rouge ⚠️', 'Lire le message : fichier manquant ou Excel ouvert ailleurs.'],
    ['Bouton PDF en échec', 'Vérifier Word installé et aucune fenêtre Word bloquante.'],
])

doc.save(SORTIE)
print(f'✅ Manuel Word généré : {SORTIE}')