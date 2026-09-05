#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_roman.py – Outil universel de mise en page KDP (v2.9.5)
v2.9.5 – OPTIMISATION VITESSE (génération 600 pages > 5 min → ~1 min) :
- instance Word UNIQUE réutilisée pour toute la génération (supprime ~45
  démarrages/fermetures coûteux de Word) ;
- mode --rapide : parité des pages estimée pendant la construction (0 comptage
  Word par bloc), comptages réels réservés à la fin (convergence + stats) ;
- suppression des sleeps/re-paginations redondants.
v2.9.4 – CORRECTIF MARGES PAIRES/IMPAIRES (rejets KDP) :
- SUPPRESSION des marges miroir (mirrorMargins) ;
- marges SYMÉTRIQUES explicites : gauche = droite = gouttière officielle
  du palier + 0,125 po de garde ;
- convergence conservée : recalage sur le nombre de pages RÉEL (Word).
v2.9.x – INTÉGRATION :
- configuration JSON prioritaire (repli Excel) ; police lettrine configurable ;
- alignements verticaux déterministes ; formats numériques ; 8.5x8.5 ; --json.
Usage :
python generer_roman.py              → génère le livre (mode exact)
python generer_roman.py --rapide     → génération rapide (parité estimée)
python generer_roman.py --init       → crée seulement l'Excel
python generer_roman.py --verif      → auto-contrôle des formats
python generer_roman.py --json       → crée/rafraîchit Configuration_roman.json
"""
import os
import re
import sys
import glob
import stat
import tempfile
import time
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# v2.9.5 : mode rapide (parité estimée pendant la construction)
MODE_RAPIDE = '--rapide' in sys.argv[1:]

# ─────────────────────────────────────────────
# 0. SOURCE DE CONFIGURATION (JSON prioritaire)
# ─────────────────────────────────────────────
try:
    import configuration_store as cs
    JSON_OK = True
except Exception:
    cs = None
    JSON_OK = False

# ─────────────────────────────────────────────
# 1. CHEMINS
# ─────────────────────────────────────────────
if getattr(sys, 'frozen', False):
    BASE = os.path.dirname(sys.executable)
else:
    BASE = os.path.dirname(os.path.abspath(__file__))

DOSSIER_IMAGES = os.path.join(BASE, 'Images')
DOSSIER_CHAPITRES = os.path.join(BASE, 'Chapitres')
DOSSIER_CACHE_HD = os.path.join(BASE, '_cache_HD')
CHEMIN_CONFIG = os.path.join(BASE, 'Configuration_roman.xlsx')
CHEMIN_CONFIG_JSON = os.path.join(BASE, 'Configuration_roman.json')

ANCIENS_INFOS = [
    os.path.join(BASE, 'Infos_roman.xlsx'),
    os.path.join(BASE, 'Infos.xlsx')
]
ANCIEN_ORGA = os.path.join(BASE, 'Organisation fichiers-chapitres.xlsx')

os.makedirs(BASE, exist_ok=True)

MOTS_TOTAL = 0
NB_ACTES = 0
NB_CHAPITRES = 0
NB_DIALOGUES = 0
NB_ILLUSTRATIONS = 0
STATS_CHAPITRES = []
SANS_TITRE = ('1.1',)

import regles as _R
FORMATS_LIVRE = _R.FORMATS_LIVRE

# ─────────────────────────────────────────────
# 2. CONFIGURATION UNIQUE (3 onglets)
# ─────────────────────────────────────────────
TITRE = 'Titre complet du roman'
SOUS_TITRE = 'Sous-titre éventuel'
AUTEUR = "Nom de l'auteur (couverture)"
PREFACE = 'Préface / Postface'
ISBN = 'Numéro ISBN'
DEPOT = 'Dépôt légal'
ANNEE = 'Année de publication'
EDITEUR = "Maison d'édition / auto-édition"
COPYRIGHT = 'Mention de copyright'
EDITION = 'Édition'
SITE = 'Site web'
AVERT = 'Avertissement'
DEDICACE = 'Dédicace'
EPIGRAPHE = 'Épigraphe'

CHAMPS_LABELS = [
    TITRE, SOUS_TITRE, AUTEUR, PREFACE, ISBN, DEPOT, ANNEE,
    EDITEUR, COPYRIGHT, EDITION, SITE, AVERT, DEDICACE, EPIGRAPHE
]

CORRESPONDANCES = {
    TITRE: ('titre complet',),
    SOUS_TITRE: ('sous-titre',),
    AUTEUR: ("nom de l'auteur",),
    PREFACE: ('préface',),
    ISBN: ('numéro isbn', 'isbn'),
    DEPOT: ('numéro de dépôt', 'dépôt légal'),
    ANNEE: ('année',),
    EDITEUR: ("maison d'édition", 'editeur', 'éditeur'),
    COPYRIGHT: ('mention de copyright', 'copyright'),
    EDITION: ('édition',),
    SITE: ('site web',),
    AVERT: ('avertissement',),
    DEDICACE: ('dédicace',),
    EPIGRAPHE: ('épigraphe',),
}

JSON_INFOS_KEYS = {
    TITRE: 'titre_complet', SOUS_TITRE: 'sous_titre', AUTEUR: 'auteur',
    PREFACE: 'preface_postface', ISBN: 'isbn', DEPOT: 'depot_legal',
    ANNEE: 'annee_publication', EDITEUR: 'maison_edition',
    COPYRIGHT: 'mention_copyright', EDITION: 'edition', SITE: 'site_web',
    AVERT: 'avertissement', DEDICACE: 'dedicace', EPIGRAPHE: 'epigraphe',
}


def nettoyer(v):
    return str(v).strip().strip('`').strip('*').strip() if v is not None else ''


def _creer_onglet_style(wb):
    from openpyxl.worksheet.datavalidation import DataValidation
    ws_s = wb.create_sheet('Style')
    ws_s.append(['Élément', 'Valeur'])
    ws_s.append(['Format du livre', '7 x 10 po'])
    ws_s.append(['Police du corps de texte', 'Aptos'])
    ws_s.append(['Taille du corps (pt)', 11])
    ws_s.append(['Police des titres', 'Cinzel'])
    ws_s.append(["Taille des titres d'acte (pt)", 16])
    ws_s.append(['Taille chapitre - ligne 1 (pt)', 14])
    ws_s.append(['Taille chapitre - ligne 2 (pt)', 13])
    ws_s.append(['Taille sous-chapitre (pt)', 12])
    ws_s.append(['Interligne du corps', 1.0])
    ws_s.column_dimensions['A'].width = 38
    ws_s.column_dimensions['B'].width = 28
    ws_l = wb.create_sheet('Listes')
    for f, _w, _h in FORMATS_LIVRE:
        ws_l.append([f])
    ws_l.sheet_state = 'hidden'
    dv = DataValidation(
        type='list',
        formula1='=Listes!$A$1:$A$' + str(len(FORMATS_LIVRE)),
        allow_blank=False
    )
    dv.error = 'Choisissez un format dans la liste.'
    dv.errorTitle = 'Format invalide'
    ws_s.add_data_validation(dv)
    dv.add(ws_s['B2'])


def creer_config_si_absent():
    if os.path.isfile(CHEMIN_CONFIG):
        return
    if JSON_OK and os.path.isfile(CHEMIN_CONFIG_JSON):
        return
    try:
        from openpyxl import Workbook, load_workbook
    except ImportError:
        print('   ⚠️  openpyxl est requis : lancez « pip install openpyxl » puis relancez.')
        raise
    wb = Workbook()
    ws_i = wb.active
    ws_i.title = 'Informations'
    valeurs = {l: '' for l in CHAMPS_LABELS}
    src = next((c for c in ANCIENS_INFOS if os.path.isfile(c)), None)
    if src:
        old = {}
        for row in load_workbook(src, data_only=True).active.iter_rows(values_only=True):
            if row and row[0] is not None:
                old[nettoyer(row[0]).lower()] = nettoyer(row[1]) if len(row) > 1 else ''
        for label in CHAMPS_LABELS:
            for prefix in CORRESPONDANCES[label]:
                for k, v in old.items():
                    if k.startswith(prefix) and v:
                        valeurs[label] = v
                        break
                if valeurs[label]:
                    break
        print('   🔁 Valeurs administratives migrées.')
    ws_i.append(['Élément', 'Votre réponse'])
    for label in CHAMPS_LABELS:
        ws_i.append([label, valeurs[label]])
    ws_i.column_dimensions['A'].width = 45
    ws_i.column_dimensions['B'].width = 70
    ws_c = wb.create_sheet('Chapitres')
    ws_c.append(['Fichier source', 'Acte', 'Chapitre Ligne 1',
                 'Chapitre Ligne 2', 'Image', 'Légende'])
    if os.path.isfile(ANCIEN_ORGA):
        for row in load_workbook(ANCIEN_ORGA, data_only=True).active.iter_rows(values_only=True):
            cells = [nettoyer(c) for c in row]
            if not any(cells) or cells[0].lower().startswith('fichier') or cells[0].startswith('---'):
                continue
            ws_c.append(cells[:6] if len(cells) >= 6 else cells + [''] * (6 - len(cells)))
        print('   🔁 Organisation des chapitres migrée.')
    for col, w in zip('ABCDEF', (38, 24, 20, 28, 14, 42)):
        ws_c.column_dimensions[col].width = w
    _creer_onglet_style(wb)
    wb.save(CHEMIN_CONFIG)
    print('   📋 Configuration_roman.xlsx créé : ' + CHEMIN_CONFIG)


def ajouter_onglet_style_si_absent():
    from openpyxl import load_workbook
    wb = load_workbook(CHEMIN_CONFIG)
    if 'Style' in wb.sheetnames:
        return
    _creer_onglet_style(wb)
    wb.save(CHEMIN_CONFIG)
    print('   🎨 Onglet Style ajouté à Configuration_roman.xlsx.')


def initialiser():
    creer_config_si_absent()
    if os.path.isfile(CHEMIN_CONFIG):
        ajouter_onglet_style_si_absent()


def _normaliser_cles(d):
    if isinstance(d, dict):
        return {str(k).strip(): _normaliser_cles(v) for k, v in d.items()}
    if isinstance(d, list):
        return [_normaliser_cles(x) for x in d]
    if isinstance(d, str):
        return d.strip()
    return d


def _infos_de(d):
    d = _normaliser_cles(d or {})
    infos = d.get('informations') if isinstance(d, dict) else None
    return infos if isinstance(infos, dict) else {}


def rafraichir_json_depuis_excel():
    if not JSON_OK:
        print('   ⚠️  Module configuration_store indisponible.')
        return False
    if not os.path.isfile(CHEMIN_CONFIG):
        print('   ⚠️  Configuration_roman.xlsx introuvable.')
        return False
    try:
        data = cs._lire_excel()
        if data:
            cs.sauvegarder_configuration(data)
            print("   ✅ Configuration_roman.json (re)créé depuis l'Excel.")
            return True
    except Exception as e:
        print('   ⚠️  Échec migration Excel → JSON : ' + str(e))
    return False

# ─────────────────────────────────────────────
# 3. LECTURE DE LA CONFIGURATION (JSON > Excel)
# ─────────────────────────────────────────────
def _num(v, defaut):
    try:
        return float(str(v).replace(',', '.'))
    except (TypeError, ValueError):
        return defaut


dimensions_pour_format = _R.dimensions_pour_format


def verifier_formats():
    ok = True
    tests = FORMATS_LIVRE + [
        ('10 x 12 cm', 10, 12),
        ('21 x 29,7 cm', 21, 29.7),
        ('140 x 210 mm', 14.0, 21.0)
    ]
    for nom, w, h in tests:
        lw, hh = dimensions_pour_format(nom)
        bon = abs(lw - w) < 0.05 and abs(hh - h) < 0.05
        ok = ok and bon
        print('   ' + ('✅' if bon else '⚠️') + ' ' + nom.ljust(16) +
              ' → ' + format(lw, '.2f') + ' x ' + format(hh, '.2f') + ' cm')
    return ok


def lire_style():
    s = {
        'format': '7 x 10 po',
        'police_corps': 'Aptos',
        'taille_corps': 11.0,
        'police_titres': 'Cinzel',
        'taille_acte': 16.0,
        'taille_chap1': 14.0,
        'taille_chap2': 13.0,
        'taille_sous': 12.0,
        'interligne': 1.0,
        'police_lettrine': 'Cinzel'
    }
    source = 'défaut'
    if JSON_OK:
        try:
            data = cs.charger_configuration()
            style = data.get('style', {})
            if style:
                s['format'] = str(style.get('format_livre') or s['format'])
                s['police_corps'] = str(style.get('police_corps') or s['police_corps'])
                s['taille_corps'] = _num(style.get('taille_corps_pt'), s['taille_corps'])
                s['police_titres'] = str(style.get('police_titres') or s['police_titres'])
                s['taille_acte'] = _num(style.get('taille_titres_acte_pt'), s['taille_acte'])
                s['taille_chap1'] = _num(style.get('taille_chapitre_ligne1_pt'), s['taille_chap1'])
                s['taille_chap2'] = _num(style.get('taille_chapitre_ligne2_pt'), s['taille_chap2'])
                s['taille_sous'] = _num(style.get('taille_sous_chapitre_pt'), s['taille_sous'])
                s['interligne'] = _num(style.get('interligne_corps'), s['interligne'])
                s['police_lettrine'] = str(style.get('police_lettrine') or s['police_titres'])
                source = 'JSON'
        except Exception:
            pass
    if source == 'défaut' and os.path.isfile(CHEMIN_CONFIG):
        from openpyxl import load_workbook
        wb = load_workbook(CHEMIN_CONFIG, data_only=True)
        if 'Style' in wb.sheetnames:
            vals = {}
            for row in wb['Style'].iter_rows(values_only=True):
                if row and row[0] is not None:
                    vals[nettoyer(row[0]).lower()] = row[1] if len(row) > 1 else None

            def get(prefix):
                for k, v in vals.items():
                    if k.startswith(prefix) and v not in (None, ''):
                        return v
                return None
            s['format'] = str(get('format') or s['format'])
            s['police_corps'] = str(get('police du corps') or s['police_corps'])
            s['taille_corps'] = _num(get('taille du corps'), 11)
            s['police_titres'] = str(get('police des titres') or s['police_titres'])
            s['taille_acte'] = _num(get("taille des titres d'acte"), 16)
            s['taille_chap1'] = _num(get('taille chapitre - ligne 1'), 14)
            s['taille_chap2'] = _num(get('taille chapitre - ligne 2'), 13)
            s['taille_sous'] = _num(get('taille sous-chapitre'), 12)
            s['interligne'] = _num(get('interligne'), 1.0)
            source = 'Excel'
    s['largeur_cm'], s['hauteur_cm'] = dimensions_pour_format(s['format'])
    s['largeur_po'] = s['largeur_cm'] / 2.54
    s['hauteur_po'] = s['hauteur_cm'] / 2.54
    s['_source'] = source
    return s


def _lire_json_brut():
    try:
        import json
        with open(CHEMIN_CONFIG_JSON, encoding='utf-8') as f:
            return _normaliser_cles(json.load(f))
    except Exception:
        return None

def _lire_json_brut():
    try:
        import json
        with open(CHEMIN_CONFIG_JSON, encoding='utf-8') as f:
            return _normaliser_cles(json.load(f))
    except Exception:
        return None


def reordonner_structure_word_file(chemin):
    """Crée les blocs manquants puis réassemble le Word dans l'ordre demandé."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches
    import json as _json, os as _os
    doc = Document(chemin)
    body = doc.element.body
    sect_modele = body.find(qn('w:sectPr'))
    inf = {}
    try:
        p = _os.path.join(BASE, 'Configuration_roman.json')
        if _os.path.isfile(p):
            inf = _infos_de(_json.load(open(p, encoding='utf-8')))
    except Exception:
        inf = {}
    def val(*ks):
        for k in ks:
            v = inf.get(k)
            if v: return str(v).strip()
        return ''
    ax = {'frontispice': val('frontispice', 'Frontispice'),
          'preface': val('preface', 'Préface'),
          'postface': val('postface', 'Postface'),
            'remerciements': val('remerciements', 'Remerciements'),
            'autres_livres': val('autres_livres', 'Autres livres du même auteur')}
    sv = inf.get('sommaire')
    ax['sommaire'] = sv if isinstance(sv, bool) else str(sv).strip().lower() != 'false'
    def a_saut(el):
        if el.tag.endswith('}p'):
            for br in el.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page': return True
            if el.find(qn('w:pPr') + '/' + qn('w:sectPr')) is not None: return True
        return False
    elems = [el for el in body if el is not sect_modele]
    blocs, cur = [], []
    for el in elems:
        cur.append(el)
        if a_saut(el): blocs.append(cur); cur = []
    if cur: blocs.append(cur)
    def texte(blk):
        return ' '.join(''.join(n.text or '' for n in el.iter(qn('w:t'))) for el in blk).strip()
    def a_image(blk):
        return any(True for _b in blk for _ in _b.iter(qn('w:drawing')))
    tit_court = re.split(r'[–-]', val('titre complet du roman', 'T'))[0].strip().upper()
    aut = val("nom de l'auteur (couverture)", 'Auteur')
    trouve, contenu = {}, []
    for blk in blocs:
        t = texte(blk); low = t.lower(); cle = None
        if 'table des matières' in low: cle = 'sommaire'
        elif '©' in t or 'isbn' in low: cle = 'copyright'
        elif val('dédicace','Dédicace') and val('dédicace','Dédicace')[:12] in t: cle = 'dedicace'
        elif val('épigraphe','Épigraphe') and val('épigraphe','Épigraphe')[:12] in t: cle = 'epigraphe'
        elif low.startswith('préface'): cle = 'preface'
        elif low.startswith('postface'): cle = 'postface'
        elif low.startswith('remerciements'): cle = 'remerciements'
        elif low.startswith('du même auteur'): cle = 'autres_livres'
        elif a_image(blk): cle = 'frontispice'
        elif tit_court and tit_court in t and '©' not in t:
            cle = 'titre' if (aut and aut.lower() in low) or len(t) > len(tit_court) + 8 else 'faux_titre'
        if cle and cle not in trouve: trouve[cle] = blk
        else: contenu.append(blk)
    # ── création des blocs manquants ──
    if 'garde' not in trouve:
        trouve['garde'] = [doc.add_paragraph()._element]
    if 'faux_titre' not in trouve and val('titre complet du roman'):
        p = doc.add_paragraph(); p.alignment = 1
        r = p.add_run(val('titre complet du roman')); r.font.size = Pt(11)
        trouve['faux_titre'] = [p._element]
    if 'frontispice' not in trouve and ax['frontispice']:
        fp = _os.path.join(BASE, 'Images', ax['frontispice'])
        if not _os.path.isfile(fp):
            for ext in ('.png', '.jpg', '.jpeg', '.webp', '.bmp'):
                if _os.path.isfile(fp + ext): fp += ext; break
        if _os.path.isfile(fp):
            try:
                doc.add_picture(fp, width=Inches(4.2))
                trouve['frontispice'] = [doc.paragraphs[-1]._element]
            except Exception: pass
    for cle, titre in (('preface', 'Préface'), ('postface', 'Postface')):
        if cle not in trouve and ax[cle]:
            pc = _os.path.join(BASE, 'Chapitres', ax[cle] + '.md')
            if _os.path.isfile(pc):
                els = []
                h = doc.add_paragraph(titre)
                try: h.style = doc.styles['Heading 1']
                except Exception: pass
                els.append(h._element)
                for ln in open(pc, encoding='utf-8').read().split('\n'):
                    ln = ln.strip()
                    if ln and not ln.startswith('# '):
                        els.append(doc.add_paragraph(ln)._element)
                trouve[cle] = els
    if 'remerciements' not in trouve and ax['remerciements']:
        els = []
        h = doc.add_paragraph('Remerciements')
        try: h.style = doc.styles['Heading 1']
        except Exception: pass
        els.append(h._element)
        for ln in ax['remerciements'].split('\n'):
            if ln.strip(): els.append(doc.add_paragraph(ln.strip())._element)
        trouve['remerciements'] = els
    if 'autres_livres' not in trouve and ax['autres_livres']:
        els = []
        h = doc.add_paragraph('Du même auteur')
        try: h.style = doc.styles['Heading 1']
        except Exception: pass
        els.append(h._element)
        for ln in ax['autres_livres'].split('\n'):
            if ln.strip(): els.append(doc.add_paragraph('• ' + ln.strip())._element)
        trouve['autres_livres'] = els
    # ── réassemblage dans l'ordre demandé ──
    for el in list(body):
        if el is not sect_modele: body.remove(el)
    def ajouter(blk, saut):
        if saut and blk:
            lp = None
            for el in reversed(list(body)):
                if el.tag.endswith('}p'): lp = el; break
            if lp is not None:
                pPr = lp.get_or_add_pPr()
                sect = deepcopy(sect_modele) if sect_modele is not None else OxmlElement('w:sectPr')
                t = sect.find(qn('w:type'))
                if t is None:
                    t = OxmlElement('w:type'); sect.insert(0, t)
                t.set(qn('w:val'), 'oddPage' if saut == 'impair' else ('evenPage' if saut == 'pair' else 'nextPage'))
                old = pPr.find(qn('w:sectPr'))
                if old is not None: pPr.remove(old)
                for r in list(lp.findall(qn('w:r'))):
                    for br in r.findall(qn('w:br')):
                        if br.get(qn('w:type')) == 'page': lp.remove(r)
                pPr.append(sect)
        for el in blk: body.append(el)
    from copy import deepcopy
    from docx.oxml import OxmlElement
    ordre = ['garde', 'faux_titre', 'frontispice', 'titre', 'copyright', 'dedicace', 'epigraphe']
    if ax['sommaire']: ordre.append('sommaire')
    ordre.append('preface')
    for cle in ordre:
        if not trouve.get(cle):
            continue
        if cle == 'frontispice':
            saut = 'pair'
        elif cle in ('titre', 'copyright'):
            saut = 'impair' if cle == 'titre' else 'pair'
        else:
            saut = 'impair'
        ajouter(trouve[cle], saut)
    for i, blk in enumerate(contenu):
        ajouter(blk, 'impair' if i == 0 else None)
    for cle in ('postface', 'remerciements', 'autres_livres'):
        if trouve.get(cle): ajouter(trouve[cle], 'impair')
    body.append(sect_modele)
    doc.save(chemin)
    return [b for b in ordre if trouve.get(b)] + ['manuscrit']


def _sauver_et_reordonner(doc, chemin):
    doc.save(chemin)
    try:
        reordonner_structure_word_file(chemin)
        print('   📚 Structure éditoriale appliquée au Word.')
    except Exception as e:
        print('   ⚠️  Réordonnancement ignoré :', e)



def lire_annexes():
    ax = {'sommaire': True}
    j = {}
    try:
        import json as _json2
        p = os.path.join(BASE, 'Configuration_roman.json')
        if os.path.isfile(p):
            j = _infos_de(_json2.load(open(p, encoding='utf-8')))
    except Exception:
        j = {}
    def g(*cles):
        for c in cles:
            v = j.get(c)
            if v: return str(v).strip()
        return ''
    ax['editeur'] = g('editeur', 'Éditeur', 'Editeur', 'maison_edition')
    ax['autres_livres'] = g('autres_livres', 'Autres livres du même auteur')
    ax['remerciements'] = g('remerciements', 'Remerciements')
    ax['frontispice'] = g('frontispice', 'Frontispice')
    ax['preface'] = g('preface', 'Préface')
    ax['postface'] = g('postface', 'Postface')
    sv = j.get('sommaire')
    ax['sommaire'] = sv if isinstance(sv, bool) else (str(sv).strip().lower() != 'false' if sv is not None else True)
    return ax

def _get_infos(d):
    d = d or {}
    if 'informations' in d: return d['informations'] or {}
    if 'informations ' in d: return d['informations '] or {}
    for k, v in d.items():
        if str(k).strip() == 'informations' and isinstance(v, dict): return v
    return {}


def _get_section(d, nom, defaut=None):
    """Recherche tolérante d'une section JSON (clés avec espaces finaux)."""
    if defaut is None: defaut = {}
    if not isinstance(d, dict): return defaut
    if nom in d and d[nom] is not None: return d[nom]
    for k, v in d.items():
        if str(k).strip() == nom and v is not None: return v
    return defaut

def _section_infos(d):
    return _get_section(d, 'informations', {}) or {}

def lire_infos():
    infos = {l: '' for l in CHAMPS_LABELS}
    source = 'défaut'
    ji = _get_infos(_lire_json_brut())
    for label in CHAMPS_LABELS:
        for k, v in ji.items():
            if v is not None and str(v).strip() and nettoyer(k).lower().startswith(CORRESPONDANCES[label]):
                infos[label] = nettoyer(v)
                source = 'JSON'
                break
    if JSON_OK:
        try:
            data = cs.charger_configuration()
            jj = _infos_de(data)
            for label, cle in JSON_INFOS_KEYS.items():
                v = jj.get(cle, '')
                if v and not infos[label]:
                    infos[label] = str(v).strip()
                    if source == 'défaut':
                        source = 'JSON'
        except Exception:
            pass
    if source == 'défaut' and os.path.isfile(CHEMIN_CONFIG):
        from openpyxl import load_workbook
        wb = load_workbook(CHEMIN_CONFIG, data_only=True)
        ws = wb['Informations'] if 'Informations' in wb.sheetnames else wb.active
        for row in ws.iter_rows(values_only=True):
            if not row or row[0] is None:
                continue
            k = nettoyer(row[0]).lower()
            if not k or k == 'élément':
                continue
            v = nettoyer(row[1]) if len(row) > 1 and row[1] is not None else ''
            for label in CHAMPS_LABELS:
                if any(k.startswith(p) for p in CORRESPONDANCES[label]):
                    infos[label] = v
                    source = 'Excel'
                    break
    infos['_source'] = source
    return infos

def lire_organisation():
    if JSON_OK:
        try:
            data = cs.charger_configuration()
            chapitres = data.get('chapitres', [])
            flux = []
            for item in chapitres:
                t = item.get('type', '')
                if t == 'image':
                    if item.get('image'):
                        flux.append({'type': 'image',
                                     'image': item.get('image', ''),
                                     'legende': item.get('legende', '')})
                elif t == 'acte':
                    if item.get('acte'):
                        flux.append({'type': 'acte', 'acte': item.get('acte', '')})
                elif t == 'chapitre':
                    if item.get('fichier_source'):
                        l1 = item.get('chapitre_ligne1', '') or ''
                        l2 = item.get('chapitre_ligne2', '') or ''
                        titre = (l1 + ' ' + l2).strip() if l2 else l1
                        flux.append({'type': 'chapitre',
                                     'fichier': item.get('fichier_source', ''),
                                     'titre': titre})
            if flux:
                return flux
        except Exception:
            pass
    if not os.path.isfile(CHEMIN_CONFIG):
        return None
    from openpyxl import load_workbook
    wb = load_workbook(CHEMIN_CONFIG, data_only=True)
    if 'Chapitres' not in wb.sheetnames:
        return None
    ws = wb['Chapitres']
    cols = {'fichier': 0, 'acte': 1, 'l1': 2, 'l2': 3, 'image': 4, 'legende': 5}
    flux = []
    for row in ws.iter_rows(values_only=True):
        cells = [nettoyer(c) for c in row]
        if not any(cells):
            continue
        low = [c.lower() for c in cells]
        if low[0].startswith('fichier'):
            for i, n in enumerate(low):
                if n.startswith('fichier'):
                    cols['fichier'] = i
                elif n.startswith('acte'):
                    cols['acte'] = i
                elif 'ligne 1' in n:
                    cols['l1'] = i
                elif 'ligne 2' in n:
                    cols['l2'] = i
                elif n.startswith('image'):
                    cols['image'] = i
                elif 'égende' in n:
                    cols['legende'] = i
            continue

        def get(col):
            idx = cols[col]
            return cells[idx] if idx < len(cells) else ''
        if get('fichier').startswith('---'):
            continue
        fichier, acte = get('fichier'), get('acte')
        l1, l2, image, legende = get('l1'), get('l2'), get('image'), get('legende')
        if image and not fichier and not acte:
            flux.append({'type': 'image', 'image': image, 'legende': legende})
        elif not fichier and acte:
            flux.append({'type': 'acte', 'acte': acte})
        elif fichier:
            titre = (l1 + ' ' + l2).strip() if l2 else l1
            flux.append({'type': 'chapitre', 'fichier': fichier, 'titre': titre})
    return flux or None

# ─────────────────────────────────────────────
# 4. CORRECTIONS & CONVERSIONS TYPOGRAPHIQUES
# ─────────────────────────────────────────────
CORRECTIONS_COMMUNES = _R.CORRECTIONS_COMMUNES
CORR_PAR_PREFIXE = _R.CORR_PAR_PREFIXE
REGEX_PENSEES = _R.REGEX_PENSEES

def corr_pour(fichier):
    return _R.corrections_pour(fichier)


# ─────────────────────────────────────────────
# 5. CHARGEMENT MARKDOWN → STRUCTURE WORD
# ─────────────────────────────────────────────
META_RE = re.compile(
    r'^(titre|œuvre|oeuvre|auteur|pages_pdf|nombre_de_pages|source|'
    r'version|d[ée]p[ôo]t_l[ée]gal|langue)\s*:',
    re.IGNORECASE
)
INLINE_RE = re.compile(r'(\*\*[^*]+?\*\*|\*[^*]+?\*|_[^_]+?_)')


def chemin_image(nom):
    if not nom:
        return None
    nom = str(nom).strip()
    if not os.path.isdir(DOSSIER_IMAGES):
        return None
    direct = os.path.join(DOSSIER_IMAGES, nom)
    if os.path.isfile(direct):
        return direct
    nom_cf = nom.casefold()
    tige_cf = os.path.splitext(nom)[0].casefold()
    fichiers = [
        f for f in os.listdir(DOSSIER_IMAGES)
        if os.path.isfile(os.path.join(DOSSIER_IMAGES, f))
    ]
    for f in fichiers:
        if f.casefold() == nom_cf:
            return os.path.join(DOSSIER_IMAGES, f)
    for f in fichiers:
        if os.path.splitext(f)[0].casefold() == tige_cf:
            return os.path.join(DOSSIER_IMAGES, f)
    for f in fichiers:
        if os.path.splitext(f)[0].casefold().startswith(tige_cf):
            return os.path.join(DOSSIER_IMAGES, f)
    return None


def image_haute_definition(chemin, largeur_cm, journal=print):
    try:
        from PIL import Image as PIL
    except Exception:
        return chemin
    try:
        os.makedirs(DOSSIER_CACHE_HD, exist_ok=True)
        cache = os.path.join(
            DOSSIER_CACHE_HD,
            os.path.splitext(os.path.basename(chemin))[0] +
            '_HD' + str(int(os.path.getmtime(chemin))) + '.png'
        )
        if not os.path.isfile(cache):
            with PIL.open(chemin) as im:
                im.load()
                pw, ph = im.size
                cible = int(round((largeur_cm / 2.54) * 300))
                if cible > pw:
                    f = cible / pw
                    resample = getattr(PIL, 'LANCZOS', getattr(PIL, 'BILINEAR', 2))
                    im = im.resize((cible, int(ph * f)), resample)
                    journal('   🔍 ' + os.path.basename(chemin) + ' : upscale HD → 300 dpi.')
                im.convert('L').save(cache, format='PNG', dpi=(300, 300))
        return cache
    except Exception:
        return chemin


def runs_riches(texte):
    runs, pos = [], 0
    for m in INLINE_RE.finditer(texte):
        if m.start() > pos:
            runs.append((texte[pos:m.start()], False, False))
        seg = m.group(0)
        if seg.startswith('**'):
            runs.append((seg[2:-2], True, False))
        else:
            runs.append((seg[1:-1], False, True))
        pos = m.end()
    if pos < len(texte):
        runs.append((texte[pos:], False, False))
    return runs or [(texte, False, False)]


def charger_chapitre(cfg):
    global MOTS_TOTAL, NB_DIALOGUES
    cand = glob.glob(os.path.join(DOSSIER_CHAPITRES, cfg['fichier'] + '*.md'))
    if not cand:
        print('   ⚠️  Fichier introuvable : ' + cfg['fichier'] + '*.md')
        return None
    with open(cand[0], encoding='utf-8') as f:
        texte = f.read()
    for a, b in CORRECTIONS_COMMUNES + cfg.get('corr', []):
        texte = texte.replace(a, b)
    while '  ' in texte:
        texte = texte.replace('  ', ' ')
    lignes = [l.strip() for l in texte.splitlines()]
    if lignes and lignes[0] == '---':
        for j in range(1, len(lignes)):
            if lignes[j] == '---':
                lignes = lignes[j + 1:]
                break
    items = []
    for l in lignes:
        if not l or META_RE.match(l):
            continue
        if cfg.get('skip_titre') and l == cfg['skip_titre']:
            continue
        if l.startswith('## '):
            items.append(('h2', l[3:].strip()))
        elif l.startswith('# '):
            items.append(('h1', l[2:].strip()))
        elif l in ('---', '***', '___', '---**---'):
            items.append(('sep', None))
        else:
            if l.startswith('—'):
                l = l.replace('« ', '').replace(' »', '')
                l = l.replace('«', '').replace('»', '')
            else:
                for rx, repl in REGEX_PENSEES:
                    l = rx.sub(repl, l)
            items.append(('p', l))
    while items and items[0][0] == 'sep':
        items.pop(0)
    if items and items[0][0] in ('h1', 'h2') and not cfg.get('sans_titre'):
        items.pop(0)
    while items and items[0][0] == 'sep':
        items.pop(0)
    paras = [t for k, t in items if k == 'p']
    mots_ch = sum(len(t.split()) for t in paras)
    dia_ch = sum(1 for t in paras if t.startswith('—'))
    MOTS_TOTAL += mots_ch
    NB_DIALOGUES += dia_ch
    STATS_CHAPITRES.append({'titre': cfg['titre'], 'mots': mots_ch})
    print('   📄 ' + os.path.basename(cand[0]) + ' → ' + cfg['titre'] +
          ' (' + str(len(items)) + ' blocs)')
    return items

# ─────────────────────────────────────────────
# 6. DOCUMENT + STYLES
# ─────────────────────────────────────────────
initialiser()
STYLE = lire_style()
POLICE_CORPS = STYLE['police_corps']
POLICE_TITRES = STYLE['police_titres']
POLICE_LETTRINE = STYLE.get('police_lettrine', STYLE['police_titres'])
TC = STYLE['taille_corps']
INTERLIGNE = STYLE['interligne']
RETRAIT_CHAPITRE = Cm(0.5)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(STYLE['largeur_cm'])
section.page_height = Cm(STYLE['hauteur_cm'])
section.top_margin = section.bottom_margin = Cm(1.9)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)
# v2.9.4 : PAS de marges miroir — marges symétriques explicites,
# seules interprétables de façon identique par Word, l'export PDF et KDP.
doc.settings.element.append(OxmlElement('w:doNotCompressImages'))
_ALIGN_VOULU = {id(section._sectPr): 'top'}

styles = doc.styles
sn = styles['Normal']
sn.font.name = POLICE_CORPS
sn.font.size = Pt(TC)
sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sn.paragraph_format.line_spacing = INTERLIGNE
sn.paragraph_format.space_before = Pt(TC)
sn.paragraph_format.space_after = Pt(0)


def creer_style(nom, police, taille, gras=False, base=None):
    s = styles.add_style(nom, 1)
    if base:
        s.base_style = styles[base]
    s.font.name = police
    s.font.size = Pt(taille)
    s.font.bold = gras
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.line_spacing = INTERLIGNE
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    return s


creer_style('TitreActe', POLICE_TITRES, STYLE['taille_acte'], True, base='Heading 1')
creer_style('TitreChapitre', POLICE_TITRES, STYLE['taille_chap1'], True, base='Heading 1')
creer_style('TitreSousChap', POLICE_CORPS, STYLE['taille_sous'], True, base='Heading 2')

sc = styles.add_style('CorpsTexte', 1)
sc.font.name = POLICE_CORPS
sc.font.size = Pt(TC)
sc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sc.paragraph_format.line_spacing = INTERLIGNE
sc.paragraph_format.space_before = Pt(TC)
sc.paragraph_format.space_after = Pt(0)

ss = styles.add_style('SeparateurScene', 1)
ss.font.name = POLICE_CORPS
ss.font.size = Pt(TC)
ss.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
ss.paragraph_format.line_spacing = INTERLIGNE
ss.paragraph_format.space_before = Pt(TC)


def ligne_vide(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = INTERLIGNE


def run_style(par, texte, police, taille, gras=False, italique=False):
    r = par.add_run(texte)
    r.font.name = police
    r.font.size = Pt(taille)
    r.font.bold = gras
    r.font.italic = italique
    return r


def champ_page(run):
    run.font.name = POLICE_CORPS
    run.font.size = Pt(10)
    run._r.append(parse_xml('<w:fldChar ' + nsdecls('w') + ' w:fldCharType="begin"/>'))
    run._r.append(parse_xml('<w:instrText ' + nsdecls('w') + ' xml:space="preserve"> PAGE </w:instrText>'))
    run._r.append(parse_xml('<w:fldChar ' + nsdecls('w') + ' w:fldCharType="end"/>'))


def definir_entete(section, texte):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    if texte:
        r = p.add_run(texte)
        r.font.name = POLICE_CORPS
        r.font.size = Pt(9)
        r.font.italic = True
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
    section.different_first_page_header_footer = True
    fp_h = section.first_page_header
    fp_h.is_linked_to_previous = False
    pf = fp_h.paragraphs[0]
    pf.clear()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not texte:
        section.different_first_page_header_footer = False


def definir_pieds(section, numerote):
    for f in (section.footer, section.first_page_footer):
        f.is_linked_to_previous = False
        p = f.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if numerote:
            champ_page(p.add_run())


_SECTPR_ORDER = [
    qn(t) for t in (
        'w:footnotePr', 'w:endnotePr', 'w:type', 'w:pgSz', 'w:pgMar', 'w:paperSrc',
        'w:pgBorders', 'w:lnNumType', 'w:pgNumType', 'w:cols', 'w:formProt', 'w:vAlign',
        'w:noEndnote', 'w:titlePg', 'w:textDirection', 'w:bidi', 'w:rtlGutter',
        'w:docGrid', 'w:printerSettings', 'w:sectPrChange'
    )
]


def set_pgnum_start(section, start):
    sectPr = section._sectPr
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:start'), str(start))
    cible = _SECTPR_ORDER.index(qn('w:pgNumType'))
    pos = len(sectPr)
    for i, child in enumerate(sectPr):
        if child.tag in _SECTPR_ORDER and _SECTPR_ORDER.index(child.tag) > cible:
            pos = i
            break
    sectPr.insert(pos, el)


def normaliser_numerotation(section_depart):
    for sec in doc.sections:
        for el in sec._sectPr.findall(qn('w:pgNumType')):
            sec._sectPr.remove(el)
    if section_depart is not None:
        set_pgnum_start(section_depart, 1)


def nouvelle_section(start=WD_SECTION_START.NEW_PAGE, mode='top'):
    sec = doc.add_section(start)
    _ALIGN_VOULU.setdefault(id(sec._sectPr), mode)
    return sec


def _set_valign(section, mode):
    mode = (mode or 'top').lower()
    try:
        from docx.enum.section import WD_SECTION_VERTICAL_ALIGNMENT
        mapping = {
            'top': WD_SECTION_VERTICAL_ALIGNMENT.TOP,
            'center': WD_SECTION_VERTICAL_ALIGNMENT.CENTER,
            'bottom': WD_SECTION_VERTICAL_ALIGNMENT.BOTTOM,
            'justify': WD_SECTION_VERTICAL_ALIGNMENT.JUSTIFY
        }
        section.vertical_alignment = mapping.get(mode, WD_SECTION_VERTICAL_ALIGNMENT.TOP)
        return
    except Exception:
        pass
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        try:
            cible = _SECTPR_ORDER.index(qn('w:vAlign'))
            pos = len(sectPr)
            for i, child in enumerate(sectPr):
                if child.tag in _SECTPR_ORDER and _SECTPR_ORDER.index(child.tag) > cible:
                    pos = i
                    break
            sectPr.insert(pos, vAlign)
        except Exception:
            sectPr.append(vAlign)
    vAlign.set(qn('w:val'), mode)


def regler_alignement_vertical(section, mode='top'):
    _ALIGN_VOULU[id(section._sectPr)] = mode
    _set_valign(section, mode)


def appliquer_alignements_finaux():
    for sec in doc.sections:
        mode = _ALIGN_VOULU.get(id(sec._sectPr), 'top')
        _set_valign(sec, mode)


def paragraphe_corps(texte, initiale_grasse=False, retrait=False):
    p = doc.add_paragraph()
    p.style = styles['CorpsTexte']
    if retrait:
        p.paragraph_format.first_line_indent = RETRAIT_CHAPITRE
    runs = runs_riches(texte)
    if initiale_grasse and runs:
        t, b, i = runs[0]
        if t:
            r0 = p.add_run(t[0])
            r0.font.name = POLICE_LETTRINE
            r0.font.size = Pt(TC + 3)
            r0.font.bold = True
            runs[0] = (t[1:], b, i)
    for t, b, i in runs:
        if not t:
            continue
        r = p.add_run(t)
        r.font.name = POLICE_CORPS
        r.font.size = Pt(TC)
        r.font.bold = b
        r.font.italic = i
    return p


def ajouter_page_acte(titre):
    ligne_vide(10)
    p = doc.add_paragraph()
    p.style = styles['TitreActe']
    run_style(p, titre, POLICE_TITRES, STYLE['taille_acte'], True)
    ligne_vide(10)


def largeur_illustration():
    sec = doc.sections[-1]
    dispo = STYLE['largeur_cm'] - sec.left_margin.cm - sec.right_margin.cm - 1.0
    return max(5.0, min(12.0, dispo))


def dimensions_image_pour_page(chemin, max_l_cm, max_h_cm):
    px = None
    try:
        from PIL import Image as PIL
        with PIL.open(chemin) as im:
            px = im.size
    except Exception:
        try:
            from docx.image.image import Image as DocxImage
            img = DocxImage.from_file(chemin)
            px = (img.px_width, img.px_height)
        except Exception:
            return None
    try:
        px_w, px_h = float(px[0]), float(px[1])
    except Exception:
        return None
    if px_w <= 0 or px_h <= 0:
        return None
    ratio = px_w / px_h
    l = float(max_l_cm)
    h = l / ratio
    if h > float(max_h_cm):
        h = float(max_h_cm)
        l = h * ratio
    return l, h


def ajouter_page_illustration(chemin, libelle, legende='',
                              saut_avant=True, centrage_vertical=False):
    global NB_ILLUSTRATIONS
    if saut_avant:
        doc.add_page_break()
    if not centrage_vertical:
        ligne_vide(6)
    if chemin and os.path.isfile(chemin):
        try:
            # ✅ CORRECTION : taille adaptée AU FORMAT DU LIVRE (pas aux marges de section)
            max_l = max(5.0, min(12.0, STYLE['largeur_cm'] - 3.0))
            max_h = max(5.0, STYLE['hauteur_cm'] - 3.8)
            dims = dimensions_image_pour_page(chemin, max_l, max_h)
            w_cm = dims[0] if dims else max_l
            chemin_hd = image_haute_definition(chemin, w_cm)
            if centrage_vertical:
                from docx.enum.text import WD_LINE_SPACING as _WLS
                _h = dims[1] if dims else w_cm
                _dispo = STYLE['hauteur_cm'] - 3.8
                _res = 1.2 if legende else 0.0
                _haut = max(0.0, (_dispo - _h - _res) / 2.0)
                _sp = doc.add_paragraph()
                _sp.paragraph_format.space_before = Pt(0)
                _sp.paragraph_format.space_after = Pt(0)
                _sp.paragraph_format.line_spacing_rule = _WLS.EXACTLY
                _sp.paragraph_format.line_spacing = Cm(_haut)
            if dims:
                doc.add_picture(chemin_hd, width=Cm(dims[0]), height=Cm(dims[1]))
            else:
                doc.add_picture(chemin_hd, width=Cm(w_cm))
            # ✅ CORRECTION : centrage horizontal garanti
            par = doc.paragraphs[-1]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(0)
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.line_spacing = 1.0
            NB_ILLUSTRATIONS += 1
            print('   🖼️  Illustration HD intégrée : ' + os.path.basename(chemin))
        except Exception as e:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_style(p, '[ Illustration – ' + libelle + ' ]', POLICE_CORPS, 10, italique=True)
            print('   ⚠️  Image illisible : ' + libelle + ' (' + str(e) + ')')
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_style(p, '[ Illustration – ' + libelle + ' ]', POLICE_CORPS, 10, italique=True)
        print('   ⚠️  Image introuvable : ' + libelle)
    if legende:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_before = Pt(6)
        pl.paragraph_format.space_after = Pt(0)
        pl.paragraph_format.line_spacing = 1.0
        run_style(pl, legende, POLICE_CORPS, 9, italique=True)
    if not centrage_vertical:
        ligne_vide(6)

def ajouter_chapitre(titre, items, sans_titre=False):
    if not sans_titre:
        if ' : ' in titre:
            label, nom = titre.split(' : ', 1)
            p = doc.add_paragraph()
            p.style = styles['TitreChapitre']
            run_style(p, label + ' :', POLICE_TITRES, STYLE['taille_chap1'], True)
            p2 = doc.add_paragraph()
            p2.style = styles['TitreChapitre']
            p2.paragraph_format.space_before = Pt(6)
            run_style(p2, nom, POLICE_TITRES, STYLE['taille_chap2'], True)
        else:
            p = doc.add_paragraph()
            p.style = styles['TitreChapitre']
            run_style(p, titre, POLICE_TITRES, STYLE['taille_chap1'], True)
        ligne_vide(2)
    premier = True
    for kind, texte in items:
        if kind == 'h1':
            if not (sans_titre and premier):
                doc.add_page_break()
            p = doc.add_paragraph()
            p.style = styles['TitreChapitre']
            run_style(p, texte, POLICE_TITRES, STYLE['taille_chap1'], True)
            ligne_vide(2)
            premier = True
        elif kind == 'h2':
            doc.add_page_break()
            p = doc.add_paragraph()
            p.style = styles['TitreSousChap']
            run_style(p, texte, POLICE_CORPS, STYLE['taille_sous'], True)
            ligne_vide(2)
            premier = True
        elif kind == 'sep':
            p = doc.add_paragraph()
            p.style = styles['SeparateurScene']
            run_style(p, '--- ✦ ---', POLICE_CORPS, TC)
        else:
            paragraphe_corps(texte, initiale_grasse=premier, retrait=premier)
            premier = False

# ─────────────────────────────────────────────
# 7. MARGES KDP (v2.9.4 : SYMÉTRIQUES, barème officiel + sécurité)
# ─────────────────────────────────────────────
gouttiere_kdp_pour = _R.gouttiere_kdp_pour
marge_kdp_pour = _R.marge_kdp_pour
MARGES_SYMETRIQUES = _R.MARGES_SYMETRIQUES

def calculer_marges_kdp(mots, nb_actes, nb_chapitres):
    pages = 120.0
    for _ in range(8):
        m = marge_kdp_pour(pages)
        largeur_utilisable = STYLE['largeur_po'] - 2 * m
        mots_par_ligne = largeur_utilisable / 0.58
        lignes_par_page = (STYLE['hauteur_po'] - 1.5) / (0.21 * INTERLIGNE)
        pages = 4 + 2 * nb_actes + nb_chapitres + mots / (mots_par_ligne * lignes_par_page)
    m = marge_kdp_pour(pages)
    return m, m, int(pages) + 1

# ─────────────────────────────────────────────
# 8. PAGES LIMINAIRES
# ─────────────────────────────────────────────
def lignes_texte(t):
    return [l.strip() for l in t.splitlines() if l.strip()]


def ajouter_page_titre(infos):
    ligne_vide(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(p, (infos[TITRE] or 'Titre du roman').upper(), POLICE_TITRES, 24, True)
    ligne_vide(2)
    if infos[SOUS_TITRE]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_style(p, infos[SOUS_TITRE], POLICE_CORPS, 14)
    ligne_vide(4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(p, infos[AUTEUR] or 'Auteur', POLICE_CORPS, 14)


def ajouter_page_copyright(infos):
    doc.add_page_break()
    ligne_vide(10)
    titre = infos[TITRE] or 'Titre du roman'
    auteur = infos[AUTEUR] or 'Auteur'
    annee = infos[ANNEE] or '2026'
    lignes = [
        (titre + ' – ' + infos[SOUS_TITRE]) if infos[SOUS_TITRE] else titre,
        '',
        infos[COPYRIGHT] or ('© ' + annee + ' ' + auteur + '. Tous droits réservés.'),
        ''
    ]
    meta = ' – '.join(x for x in [infos[EDITION], infos[ANNEE]] if x)
    if meta:
        lignes.append(meta)
    if infos[ISBN]:
        lignes.append('ISBN : ' + infos[ISBN])
    if infos[DEPOT]:
        lignes.append('Dépôt légal : ' + infos[DEPOT])
    if infos[EDITEUR]:
        lignes.append('')
        lignes.append(infos[EDITEUR])
    lignes += [
        '',
        'Toute reproduction, même partielle, est interdite sans l’autorisation',
        'préalable de l’auteur, conformément aux dispositions de la législation',
        'en vigueur sur la propriété intellectuelle.'
    ]
    if infos[SITE]:
        lignes += ['', infos[SITE]]
    for t in lignes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(2)
        run_style(p, t, POLICE_CORPS, 9)


def ajouter_page_avertissement(infos):
    if not infos[AVERT]:
        return
    doc.add_page_break()
    ligne_vide(8)
    for t in lignes_texte(infos[AVERT]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_style(p, t, POLICE_CORPS, 10, italique=True)
    ligne_vide(8)


def ajouter_page_dedicace(infos):
    if not infos[DEDICACE]:
        return
    doc.add_page_break()
    ligne_vide(8)
    for t in lignes_texte(infos[DEDICACE]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(4)
        run_style(p, t, POLICE_CORPS, 12, italique=True)
    ligne_vide(8)


def ajouter_page_epigraphe(infos):
    if not infos[EPIGRAPHE]:
        return
    doc.add_page_break()
    ligne_vide(8)
    for t in lignes_texte(infos[EPIGRAPHE]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(4)
        run_style(p, t, POLICE_CORPS, 12, italique=True)
    ligne_vide(8)


def ajouter_page_faux_titre(infos):
    ligne_vide(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(p, infos[TITRE] or 'Titre du roman', POLICE_TITRES, 16, True)


def chemin_frontispice(nom):
    if not nom:
        return None
    chemin = os.path.join(DOSSIER_IMAGES, nom)
    if os.path.isfile(chemin):
        return chemin
    for extension in ('.png', '.jpg', '.jpeg', '.webp', '.bmp'):
        candidat = chemin + extension
        if os.path.isfile(candidat):
            return candidat
    return None


def ajouter_page_frontispice():
    chemin = chemin_frontispice(lire_annexes().get('frontispice'))
    if chemin:
        ajouter_page_illustration(
            chemin,
            'Frontispice',
            saut_avant=False,
            centrage_vertical=True,
        )

# ─────────────────────────────────────────────
# 9. TABLE DES MATIÈRES
# ─────────────────────────────────────────────
def ajouter_table_des_matieres_fin(document, titre="Table des matières"):
    pt = document.add_paragraph()
    pt.style = "Title"
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(titre)
    rt.font.name = POLICE_TITRES
    rt.font.size = Pt(STYLE['taille_chap1'])
    rt.font.bold = True
    p = document.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    run._r.append(b)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(it)
    s = OxmlElement("w:fldChar")
    s.set(qn("w:fldCharType"), "separate")
    run._r.append(s)
    ph = p.add_run("Cliquez avec le bouton droit ici puis choisissez « Mettre à jour les champs ».")
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    ph._r.append(e)


def forcer_mise_a_jour_des_champs(document):
    u = OxmlElement("w:updateFields")
    u.set(qn("w:val"), "true")
    document.settings.element.append(u)

# ─────────────────────────────────────────────
# 10. PAGES RÉELLES & STATISTIQUES
# ─────────────────────────────────────────────
# v2.9.5 : instance Word UNIQUE réutilisée (évite ~45 démarrages coûteux)
_WORD_APP = {'app': None}


def _word_app():
    if _WORD_APP['app'] is None:
        import win32com.client as win32
        app = win32.Dispatch('Word.Application')
        app.Visible = False
        try:
            app.DisplayAlerts = 0
            app.Options.Pagination = True
        except Exception:
            pass
        _WORD_APP['app'] = app
    return _WORD_APP['app']


def fermer_word():
    if _WORD_APP['app'] is not None:
        try:
            _WORD_APP['app'].Quit()
        except Exception:
            pass
        _WORD_APP['app'] = None


def chemin_temporaire_docx(prefix='kdp_comptage_'):
    fd, chemin = tempfile.mkstemp(prefix=prefix, suffix='.docx')
    os.close(fd)
    return chemin


def supprimer_fichier_securise(chemin, tentatives=8, delai=0.25):
    if not chemin:
        return False
    for _ in range(tentatives):
        try:
            if os.path.exists(chemin):
                try:
                    os.chmod(chemin, stat.S_IWRITE)
                except OSError:
                    pass
                os.remove(chemin)
                return True
        except PermissionError:
            time.sleep(delai)
        except OSError:
            time.sleep(delai)
    return False


def enregistrer_docx_securise(document, chemin, tentatives=5, delai=0.4):
    last_error = None
    for _ in range(tentatives):
        try:
            document.save(chemin)
            return chemin
        except PermissionError as e:
            last_error = e
            time.sleep(delai)
        except OSError as e:
            last_error = e
            time.sleep(delai)
    base, ext = os.path.splitext(chemin)
    alt = base + '_' + str(int(time.time())) + ext
    try:
        document.save(alt)
        print('   ⚠️  Le fichier final était verrouillé ou protégé.')
        print('   ℹ️  Le document a été enregistré sous : ' + alt)
        return alt
    except Exception:
        if last_error:
            raise last_error
        raise


def compter_pages_reelles(nom_fichier):
    """v2.9.5 : Word partagé, sans démarrages répétés ni sleeps inutiles."""
    try:
        import win32com.client  # noqa
    except ImportError:
        print('   ℹ️  pywin32 absent (pip install pywin32) → pages réelles non calculées')
        return None
    d = None
    chemin = os.path.abspath(nom_fichier)
    try:
        word = _word_app()
        d = word.Documents.Open(chemin, False, True, False)
        pages = d.ComputeStatistics(2)   # force la pagination automatiquement
        d.Close(False)
        d = None
        return pages
    except Exception as e:
        print('   ⚠️  Comptage Word impossible : ' + str(e))
        return None
    finally:
        if d is not None:
            try:
                d.Close(False)
            except Exception:
                pass


def maj_champs_word(nom_fichier):
    try:
        import win32com.client  # noqa
    except ImportError:
        print('   ℹ️  pywin32 absent : ouvrez le fichier dans Word puis Ctrl+A, F9.')
        return
    d = None
    try:
        word = _word_app()
        d = word.Documents.Open(os.path.abspath(nom_fichier), False, False, False)
        for f in d.Fields:
            try:
                f.Update()
            except Exception:
                pass
        d.Save()
        pages = d.ComputeStatistics(2)
        d.Close(False)
        d = None
        print('   ✅ TDM et champs mis à jour dans le fichier final.')
        return pages
    except Exception as e:
        print('   ⚠️  Mise à jour Word impossible : ' + str(e))
    finally:
        if d is not None:
            try:
                d.Close(False)
            except Exception:
                pass


def fmt(n):
    return format(n, ',').replace(',', ' ')


def afficher_statistiques(pages_reelles, titre_aff):
    total_min = int(round(MOTS_TOTAL / 230))
    h, mn = divmod(total_min, 60)
    lecture = (str(h) + ' h ' + str(mn).zfill(2)) if h else (str(mn) + ' min')
    mots_moyens = int(round(MOTS_TOTAL / NB_CHAPITRES)) if NB_CHAPITRES else 0
    bar = '=' * 58
    print()
    print(bar)
    print('📊 STATISTIQUES DU LIVRE')
    print(bar)
    print('   📖 Ouvrage            : ' + titre_aff)
    print('   📐 Format             : ' + STYLE['format'] + ' (' +
          format(STYLE['largeur_cm'], '.1f') + ' × ' +
          format(STYLE['hauteur_cm'], '.1f') + ' cm)')
    print('   🎭 Actes               : ' + str(NB_ACTES))
    print('   📚 Chapitres           : ' + str(NB_CHAPITRES))
    print('   💬 Dialogues (—)       : ' + fmt(NB_DIALOGUES))
    print('   🖼️  Illustrations      : ' + str(NB_ILLUSTRATIONS))
    print('   🔤 Mots                : ' + fmt(MOTS_TOTAL))
    print('   📊 Mots moy./chapitre  : ' + fmt(mots_moyens))
    print('   📑 Pages réelles       : ' + (fmt(pages_reelles) if pages_reelles else '—'))
    print('   ⏱️  Lecture estimée    : ' + lecture + ' (à 230 mots/min)')
    print('-' * 58)
    print('   Nombre de mots par chapitre :')
    for ch in STATS_CHAPITRES:
        print('   • ' + ch['titre'].ljust(40) + fmt(ch['mots']).rjust(8) + ' mots')
    print(bar)

# ─────────────────────────────────────────────
# 11. GÉNÉRATION
# ─────────────────────────────────────────────
def slug(s):
    s = re.sub(r'[^\w]+', '_', s.strip(), flags=re.UNICODE)
    return s.strip('_') or 'roman'


def main():
    global NB_ACTES, NB_CHAPITRES
    if JSON_OK and os.path.isfile(CHEMIN_CONFIG_JSON):
        print('   🧩 Source de configuration : Configuration_roman.json')
    elif os.path.isfile(CHEMIN_CONFIG):
        print('   📋 Source de configuration : Configuration_roman.xlsx (repli)')
    else:
        print('   ⚠️  Aucune configuration trouvée — valeurs par défaut.')
    if MODE_RAPIDE:
        print('   ⚡ Mode RAPIDE activé : parité estimée pendant la construction.')

    infos = lire_infos()
    titre_aff = infos[TITRE] or 'Titre du roman'
    if infos[SOUS_TITRE]:
        titre_aff += ' – ' + infos[SOUS_TITRE]
    print('   🎨 Style : ' + STYLE['format'] + ' · corps ' + POLICE_CORPS + ' ' +
          format(TC, 'g') + ' pt · titres ' + POLICE_TITRES + ' · lettrine ' +
          POLICE_LETTRINE + ' · interligne ' + format(INTERLIGNE, 'g') +
          ' (source : ' + STYLE.get('_source', '?') + ')')

    doc.add_page_break()
    doc.add_page_break()
    ajouter_page_faux_titre(infos)
    doc.add_page_break()
    ajouter_page_frontispice()
    doc.add_page_break()
    ajouter_page_titre(infos)
    ajouter_page_copyright(infos)
    ajouter_page_avertissement(infos)
    ajouter_page_dedicace(infos)
    ajouter_page_epigraphe(infos)

    organisation = lire_organisation()
    section_depart_num = None
    charges = {}
    if organisation:
        NB_ACTES = sum(1 for b in organisation if b.get('type') == 'acte')
        for bloc in organisation:
            if bloc.get('type') != 'chapitre':
                continue
            m = re.match(r'(\d+\.\d+)', bloc['fichier'])
            prefixe = m.group(1) if m else bloc['fichier'][:3]
            sans_titre = (prefixe in SANS_TITRE)
            cfg = {
                'fichier': prefixe,
                'titre': bloc['titre'],
                'corr': corr_pour(bloc['fichier']),
                'skip_titre': bloc['titre'],
                'sans_titre': sans_titre
            }
            charges[bloc['fichier']] = (prefixe, sans_titre, charger_chapitre(cfg))
        NB_CHAPITRES = sum(1 for _p, _s, it in charges.values() if it)

    def marges_courantes():
        g, o, est = calculer_marges_kdp(MOTS_TOTAL, NB_ACTES, NB_CHAPITRES)
        for s in doc.sections:
            s.left_margin = Inches(g)
            s.right_margin = Inches(g)   # v2.9.4 : symétrique
            s.top_margin = s.bottom_margin = Cm(1.9)
            s.header_distance = Cm(1.0)
            s.footer_distance = Cm(1.0)
        return est

    gutter, outside, _ = calculer_marges_kdp(MOTS_TOTAL, NB_ACTES, NB_CHAPITRES)
    for s in doc.sections:
        s.left_margin = Inches(gutter)
        s.right_margin = Inches(gutter)
    print('   📏 Marges KDP symétriques : ' + format(gutter, '.3f') +
          ' po des deux côtés (gouttière officielle + garde)')

    def pages_courantes(est):
        temp = chemin_temporaire_docx('kdp_comptage_')
        try:
            _sauver_et_reordonner(doc, temp)
            n = compter_pages_reelles(temp)
            return n if n is not None else est
        except OSError as e:
            print('   ⚠️  Impossible d’utiliser le fichier temporaire de comptage.')
            print('   ⚠️  Détail : ' + str(e))
            print('   ℹ️  Le script continue avec l’estimation de pages.')
            return est
        finally:
            supprimer_fichier_securise(temp)

    def inserer_blanks(n):
        for _ in range(n):
            sb = nouvelle_section(WD_SECTION_START.NEW_PAGE, 'top')
            definir_entete(sb, '')
            definir_pieds(sb, False)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    def blanks_pour_impaire(pages_avant):
        try:
            pages_avant = int(pages_avant)
        except Exception:
            return 0
        if pages_avant <= 0:
            return 0
        return 1 if pages_avant % 2 == 1 else 0

    def saut_vers_page_impaire(libelle):
        est = marges_courantes()
        if MODE_RAPIDE:
            pages_avant = est          # v2.9.5 : parité estimée, 0 comptage Word
        else:
            pages_avant = pages_courantes(est)
        nb = blanks_pour_impaire(pages_avant)
        inserer_blanks(nb)
        if nb:
            print('   📄 ' + str(nb) + ' page blanche avant ' + libelle + '.')
        return nouvelle_section(WD_SECTION_START.NEW_PAGE, 'top')

    print('   📊 Organisation lue depuis la configuration :')
    numero_acte = 0
    for bloc in organisation:
        if bloc.get('type') == 'image':
            sec_img = saut_vers_page_impaire('l’illustration ' + bloc['image'])
            definir_entete(sec_img, '')
            definir_pieds(sec_img, False)
            regler_alignement_vertical(sec_img, 'center')
            ajouter_page_illustration(
                chemin_image(bloc['image']),
                bloc['image'],
                legende=bloc.get('legende', ''),
                saut_avant=False,
                centrage_vertical=True
            )
            continue
        if bloc.get('type') == 'acte':
            numero_acte += 1
            sec = saut_vers_page_impaire('l’acte ' + str(numero_acte))
            definir_entete(sec, '')
            definir_pieds(sec, False)
            regler_alignement_vertical(sec, 'top')
            ajouter_page_acte(bloc['acte'])
            continue
        if bloc.get('type') == 'chapitre':
            prefixe, sans_titre, items = charges.get(bloc['fichier'], (None, False, None))
            if items:
                sec_ch = saut_vers_page_impaire(bloc['titre'])
                definir_entete(sec_ch, bloc['titre'])
                regler_alignement_vertical(sec_ch, 'top')
                if section_depart_num is None:
                    section_depart_num = sec_ch
                definir_pieds(sec_ch, True)
                ajouter_chapitre(bloc['titre'], items, sans_titre=sans_titre)
            else:
                print('   ⚠️  Configuration introuvable ou vide.')

    def marges_courantes_fin():
        g, o, est = calculer_marges_kdp(MOTS_TOTAL, NB_ACTES, NB_CHAPITRES)
        for s in doc.sections:
            s.left_margin = Inches(g)
            s.right_margin = Inches(g)
            s.top_margin = s.bottom_margin = Cm(1.9)
            s.header_distance = Cm(1.0)
            s.footer_distance = Cm(1.0)
        return est

    def pages_courantes_fin(est):
        temp = chemin_temporaire_docx('kdp_comptage_')
        try:
            _sauver_et_reordonner(doc, temp)
            n = compter_pages_reelles(temp)
            return n if n is not None else est
        except OSError:
            return est
        finally:
            supprimer_fichier_securise(temp)

    def inserer_blanks_fin(n):
        for _ in range(n):
            sb = nouvelle_section(WD_SECTION_START.NEW_PAGE, 'top')
            definir_entete(sb, '')
            definir_pieds(sb, False)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    est = marges_courantes_fin()
    pages_avant = est if MODE_RAPIDE else pages_courantes_fin(est)
    nb_blanks = 1 if (int(pages_avant) % 2 == 1 and pages_avant > 0) else 0
    inserer_blanks_fin(nb_blanks)
    if nb_blanks:
        print('   📄 ' + str(nb_blanks) + ' page blanche avant la table des matières.')

    sec_toc = nouvelle_section(WD_SECTION_START.NEW_PAGE, 'top')
    definir_entete(sec_toc, '')
    definir_pieds(sec_toc, False)
    regler_alignement_vertical(sec_toc, 'top')
    ajouter_table_des_matieres_fin(doc)
    forcer_mise_a_jour_des_champs(doc)
    normaliser_numerotation(section_depart_num)
    appliquer_alignements_finaux()

    gutter, outside, pages_est = calculer_marges_kdp(MOTS_TOTAL, NB_ACTES, NB_CHAPITRES)
    for sec in doc.sections:
        sec.left_margin = Inches(gutter)
        sec.right_margin = Inches(gutter)

    nom = slug(infos[TITRE] or 'roman')
    if infos[SOUS_TITRE]:
        nom += '_' + slug(infos[SOUS_TITRE])
    _doss = os.path.join(BASE, 'export'); os.makedirs(_doss, exist_ok=True)
    nom_fichier = os.path.join(_doss, nom + '_KDP.docx')
    doc.core_properties.title = infos[TITRE] or ''
    doc.core_properties.subject = infos[SOUS_TITRE] or ''
    doc.core_properties.author = infos[AUTEUR] or ''
    doc.core_properties.keywords = 'KDP, roman, Danoë Studio'
    nom_fichier = enregistrer_docx_securise(doc, nom_fichier)

    # ── v2.9.4 : convergence des marges sur le nombre de pages RÉEL ──
    print('✅ Document généré avec succès !')
    print('   📁 ' + nom_fichier)
    print('   📑 Comptage des pages réelles via Word…')
    pages_reelles = maj_champs_word(nom_fichier)
    if not pages_reelles:
        pages_reelles = compter_pages_reelles(nom_fichier)
    if pages_reelles:
        m_cible = marge_kdp_pour(pages_reelles)
        if abs(m_cible - gutter) >= 0.001:
            gutter = m_cible
            for sec in doc.sections:
                sec.left_margin = Inches(gutter)
                sec.right_margin = Inches(gutter)
            nom_fichier = enregistrer_docx_securise(doc, nom_fichier)
            pages_reelles = maj_champs_word(nom_fichier) or pages_reelles
        # ── CORRECTIF annexes Word : postface / remerciements / autres livres ──
    try:
        _ax = lire_annexes()
        def _titre_h(t):
            p = doc.add_paragraph(t)
            try: p.style = doc.styles['Heading 1']
            except Exception: pass
            return p
        if _ax.get('postface'):
            _pc = os.path.join(BASE, 'Chapitres', _ax['postface'] + '.md')
            if os.path.isfile(_pc):
                doc.add_page_break(); _titre_h('Postface')
                for _ln in open(_pc, encoding='utf-8').read().split('\n'):
                    _ln = _ln.strip()
                    if _ln and not _ln.startswith('# '): doc.add_paragraph(_ln)
        if _ax.get('remerciements'):
            doc.add_page_break(); _titre_h('Remerciements')
            for _ln in _ax['remerciements'].split('\n'):
                if _ln.strip(): doc.add_paragraph(_ln.strip())
        if _ax.get('autres_livres'):
            doc.add_page_break(); _titre_h('Du même auteur')
            for _ln in _ax['autres_livres'].split('\n'):
                if _ln.strip(): doc.add_paragraph('• ' + _ln.strip())
        if not _ax.get('sommaire', True):
            for _p in list(doc.paragraphs):
                if 'Table des matières' in _p.text:
                    _p.clear()
    except Exception:
        pass
    try:
        nom_fichier = enregistrer_docx_securise(doc, nom_fichier)
        reordonner_structure_word_file(nom_fichier)
        print('   📚 Structure éditoriale appliquée au Word.')
    except Exception as e:
        print('   ⚠️ Réordonnancement Word ignoré :', e)
    afficher_statistiques(pages_reelles, titre_aff)
    # ── CORRECTIF registre : régénère registre.json pour l'interface ──
    try:
        import json as _json_reg
        _r = lire_infos()
        _inf = _r[0] if isinstance(_r, tuple) else _r
        _sty = lire_style()
        _org = lire_organisation()
        _mots = _nchap = _nill = 0
        for _b in (_org or []):
            if str(_b.get('type', '')) == 'chapitre':
                _nchap += 1
                try:
                    _txt = open(os.path.join(BASE, 'Chapitres', str(_b.get('fichier_source', '')) + '.md'), encoding='utf-8').read()
                    _mots += len(_txt.split())
                    _nill += _txt.count('![')
                except Exception:
                    pass
            elif str(_b.get('type', '')) == 'image':
                _nill += 1
        _reg = {
            'ouvrage': (_inf.get('titre complet du roman', '') or titre_aff),
            'format': (_sty.get('format_livre') or _sty.get('format') or ''),
            'mots': _mots,
            'pages': int(pages_reelles or 0),
            'chapitres': _nchap,
            'illustrations': _nill,
            'source': 'JSON',
        }
        for _chem_reg in (os.path.join(BASE, 'registre.json'),
                          os.path.join(BASE, 'statistiques.json'),
                          os.path.join(BASE, 'export', 'registre.json'),
                          os.path.join(BASE, 'export', 'statistiques.json')):
            try:
                os.makedirs(os.path.dirname(_chem_reg), exist_ok=True)
                with open(_chem_reg, 'w', encoding='utf-8') as _f:
                    _json_reg.dump(_reg, _f, ensure_ascii=False, indent=2)
            except Exception:
                pass
    except Exception:
        pass



if __name__ == '__main__':
    if '--init' in sys.argv[1:]:
        initialiser()
        print('✅ Fichier de configuration prêt (3 onglets).')
    elif '--verif' in sys.argv[1:]:
        print('🔎 Contrôle des formats :')
        verifier_formats()
    elif '--json' in sys.argv[1:]:
        print('🧩 Création/rafraîchissement du fichier JSON…')
        rafraichir_json_depuis_excel()
    else:
        try:
            main()
        finally:
            fermer_word()   # v2.9.5 : ferme l'instance Word unique
# ── CORRECTIF registre v2 : clés/valeurs propres pour l'interface ──
def _normaliser_registre():
    import json as _j
    for _p in (os.path.join(BASE, 'registre.json'), os.path.join(BASE, 'export', 'registre.json')):
        try:
            if not os.path.isfile(_p):
                continue
            _d = _j.load(open(_p, encoding='utf-8'))
            _c = {str(k).strip(): (str(v).strip() if isinstance(v, str) else v) for k, v in _d.items()}
            if not _c.get('mots'):
                _tot = 0
                _dch = os.path.join(BASE, 'Chapitres')
                if os.path.isdir(_dch):
                    for _f in os.listdir(_dch):
                        if _f.lower().endswith('.md'):
                            try:
                                _tot += len(open(os.path.join(_dch, _f), encoding='utf-8').read().split())
                            except Exception:
                                pass
                _c['mots'] = _tot
            _j.dump(_c, open(_p, 'w', encoding='utf-8'), ensure_ascii=False, indent=2)
        except Exception:
            pass
import atexit as _atx_reg
_atx_reg.register(_normaliser_registre)

